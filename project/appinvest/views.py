import datetime

from django.contrib.auth.decorators import login_required
from django.http import HttpResponse, Http404
from django.shortcuts import render, redirect
from django.views.decorators.csrf import csrf_exempt
from . import models
from .models import Result
import docx


def chunk_based_on_size(lst, n):
    for x in range(0, len(lst), n):
        each_chunk = lst[x: n+x]

        if len(each_chunk) < n:
            each_chunk = each_chunk + [None for y in range(n-len(each_chunk))]
        yield each_chunk


@login_required
@csrf_exempt
def appdistrib(request):
    if request.method == 'GET':
        return render(request, 'appdistrib.html')

    elif request.method == 'POST':
        result = request.POST.get('result')
        table = request.POST.get('out_arr')
        print(result)
        print(table)
        newrec = models.Result(
            title='Решение от ' + '.'.join(str(datetime.datetime.now()).split('.')[:-1]),
            result=result,
            table=table,
            username=str(request.user.username)
        )
        newrec.save()
        return redirect('/distribution/history')
    else:
        return HttpResponse('Что-то пошло не так..')


@login_required
def history(request):
    latest_text_list = Result.objects.filter(username=str(request.user.username)).order_by('-id')
    return render(request, 'history.html', {'latest_text_list': latest_text_list})


@csrf_exempt
@login_required
def detail(request, id):
    if request.method == 'GET':
        try:
            a = Result.objects.get(id=id)
        except:
            raise Http404('Запись не найдена!')
        if a.username == request.user.username:
            result_list = a.result.split()
            table_list = a.table.split()
            Predpriyat = result_list[1]
            name_list = []
            distib_list = []
            for i in range(int(Predpriyat)):
                name_list.append(result_list[10 + (i * 4)])
                distib_list.append(result_list[12 + (i * 4)])
            new_table_list = list(chunk_based_on_size(table_list,int(Predpriyat)))
            print(new_table_list)
            return render(request, 'detail.html', {'new_table_list': new_table_list,'text': a, 'name_list': name_list, 'distib_list': distib_list})
        else:
            return HttpResponse('Запись не найдена')

    elif (request.method == "POST") and ('Удалить' in request.POST):
        a = Result.objects.get(id=id)
        a.delete()
        return redirect('/distribution/history')

    elif (request.method == "POST") and ('Сохранить в Word' in request.POST):
        a = Result.objects.get(id=id)
        doc = docx.Document()
        doc.add_paragraph(a.result)
        doc.add_paragraph('Таблица:')
        doc.add_paragraph(a.table)
        doc.save('media/Result Files/' + a.title + '_' + str(a.id) + '.docx')
        return redirect('/media/Result Files/' + a.title + '_' + str(a.id) + '.docx')

